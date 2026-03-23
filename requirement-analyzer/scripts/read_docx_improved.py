# -*- coding: utf-8 -*-
"""
文件名：read_docx_improved.py
功能：改进版Word文档读取器 - 返回原始表格数据
改进点：
1. 返回原始表格数据，使用key-value格式
2. 不做任何数据处理，保持原始内容
3. 支持多表格读取
4. 支持提取表格中的图片并保存
日期：2026-02-11
作者：刘丰荣
"""

from docx import Document
from docx.oxml.ns import qn
import sys
import json
import io
import re
import os
from typing import List, Dict, Any, Optional, Tuple

# 设置stdout编码为utf-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')


def clean_text(text: str) -> str:
    """清理文本中的特殊字符"""
    if not text:
        return ""

    # Unicode空格和特殊字符替换
    replacements = {
        '\u2002': ' ',  # EN SPACE
        '\u2003': ' ',  # EM SPACE
        '\u00a0': ' ',  # NO-BREAK SPACE
        '\u3000': ' ',  # IDEOGRAPHIC SPACE
        '\u2610': '[ ]',  # BALLOT BOX
        '\u2611': '[x]',  # BALLOT BOX WITH CHECK
        '\u2612': '[X]',  # BALLOT BOX WITH X
        '\t': ' ',      # TAB
        '\r': '',       # CARRIAGE RETURN
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    # 合并多个空格
    text = re.sub(r'\s+', ' ', text)

    return text.strip()


def has_image(cell) -> bool:
    """
    检查单元格是否包含图片

    Args:
        cell: 表格单元格对象

    Returns:
        bool: 是否包含图片
    """
    # 查找所有的blip元素（图片引用）
    blips = cell._tc.xpath('.//a:blip')
    return len(blips) > 0


# 全局变量用于跨单元格去重图片
_image_rId_map = {}  # rId -> filename

def extract_images_from_paragraph(para, doc_part, output_dir: str, para_idx: int) -> List[str]:
    """
    从段落中提取图片并保存

    Args:
        para: 段落对象
        doc_part: 文档part对象
        output_dir: 图片保存目录
        para_idx: 段落索引

    Returns:
        List[str]: 保存的图片文件名列表
    """
    global _image_rId_map
    saved_images = []

    # 查找段落中所有的blip元素（图片引用）
    blips = para._element.xpath('.//a:blip')

    for blip in blips:
        # 获取图片的关系ID
        embed_attr = qn('r:embed')
        rId = blip.get(embed_attr)

        if rId and rId in doc_part.related_parts:
            try:
                # 检查是否已经提取过该图片（按rId去重）
                if rId in _image_rId_map:
                    saved_images.append(_image_rId_map[rId])
                    continue

                image_part = doc_part.related_parts[rId]
                image_data = image_part.blob

                # 确定图片格式
                content_type = getattr(image_part, 'content_type', '')
                if 'png' in content_type:
                    ext = 'png'
                elif 'jpeg' in content_type or 'jpg' in content_type:
                    ext = 'jpg'
                elif 'gif' in content_type:
                    ext = 'gif'
                elif 'bmp' in content_type:
                    ext = 'bmp'
                else:
                    ext = 'png'  # 默认png

                # 生成文件名: para{para_idx}_img{全局序号}.{ext}
                file_name = f"para{para_idx}_img{len(_image_rId_map) + 1}.{ext}"
                file_path = os.path.join(output_dir, file_name)

                # 保存图片
                with open(file_path, 'wb') as f:
                    f.write(image_data)

                # 记录rId和文件名的映射
                _image_rId_map[rId] = file_name
                saved_images.append(file_name)

            except Exception as e:
                print(f"警告: 提取段落图片失败 - {e}", file=sys.stderr)

    return saved_images


def has_image_in_paragraph(para) -> bool:
    """
    检查段落是否包含图片

    Args:
        para: 段落对象

    Returns:
        bool: 是否包含图片
    """
    blips = para._element.xpath('.//a:blip')
    return len(blips) > 0


def extract_images_from_cell(cell, doc_part, output_dir: str, table_idx: int, row_idx: int, col_idx: int) -> List[str]:
    """
    从单元格中提取图片并保存
    使用rId进行全局去重，避免合并单元格导致重复提取

    Args:
        cell: 表格单元格对象
        doc_part: 文档part对象
        output_dir: 图片保存目录
        table_idx: 表格索引
        row_idx: 行索引
        col_idx: 列索引

    Returns:
        List[str]: 保存的图片文件名列表
    """
    global _image_rId_map
    saved_images = []
    image_count = 0

    # 查找所有的blip元素
    blips = cell._tc.xpath('.//a:blip')

    for blip in blips:
        # 获取图片的关系ID
        embed_attr = qn('r:embed')
        rId = blip.get(embed_attr)

        if rId and rId in doc_part.related_parts:
            try:
                # 检查是否已经提取过该图片（按rId去重）
                if rId in _image_rId_map:
                    saved_images.append(_image_rId_map[rId])
                    continue

                image_part = doc_part.related_parts[rId]
                image_data = image_part.blob

                # 确定图片格式
                content_type = getattr(image_part, 'content_type', '')
                if 'png' in content_type:
                    ext = 'png'
                elif 'jpeg' in content_type or 'jpg' in content_type:
                    ext = 'jpg'
                elif 'gif' in content_type:
                    ext = 'gif'
                elif 'bmp' in content_type:
                    ext = 'bmp'
                else:
                    ext = 'png'  # 默认png

                # 生成文件名: 表{table_idx}_行{row_idx}_图{全局序号}.{ext}
                image_count += 1
                file_name = f"table{table_idx}_row{row_idx}_img{len(_image_rId_map) + 1}.{ext}"
                file_path = os.path.join(output_dir, file_name)

                # 保存图片
                with open(file_path, 'wb') as f:
                    f.write(image_data)

                # 记录rId和文件名的映射
                _image_rId_map[rId] = file_name
                saved_images.append(file_name)

            except Exception as e:
                print(f"警告: 提取图片失败 - {e}", file=sys.stderr)

    return saved_images


def extract_cell_content(cell, doc_part, output_dir: str, table_idx: int, row_idx: int, col_idx: int) -> Tuple[str, List[str]]:
    """
    提取单元格内容（包括文本和图片）

    Args:
        cell: 表格单元格对象
        doc_part: 文档part对象
        output_dir: 图片保存目录
        table_idx: 表格索引
        row_idx: 行索引
        col_idx: 列索引

    Returns:
        Tuple[str, List[str]]: (文本内容, 图片文件名列表)
    """
    # 提取文本
    text = clean_text(cell.text)

    # 提取图片
    images = []
    if has_image(cell):
        images = extract_images_from_cell(cell, doc_part, output_dir, table_idx, row_idx, col_idx)

    return text, images


def extract_unique_from_row_with_images(row, doc_part, output_dir: str, table_idx: int, row_idx: int) -> Tuple[List[str], Dict[int, List[str]]]:
    """
    从表格行中提取唯一内容（包括文本和图片）
    处理合并单元格：文本和图片都去重，避免重复提取

    Args:
        row: 表格行对象
        doc_part: 文档part对象
        output_dir: 图片保存目录
        table_idx: 表格索引
        row_idx: 行索引

    Returns:
        Tuple[List[str], Dict[int, List[str]]]: (文本列表, 列索引到图片列表的映射)
    """
    seen_texts = set()
    seen_images = set()  # 用于去重图片
    unique_cells = []
    cell_images = {}  # 记录每个列的图片

    for col_idx, cell in enumerate(row.cells):
        text, images = extract_cell_content(cell, doc_part, output_dir, table_idx, row_idx, col_idx)

        # 处理文本去重
        if text and text not in seen_texts:
            seen_texts.add(text)
            unique_cells.append(text)

        # 处理图片去重（合并单元格会导致同一个图片出现在多个cell中）
        if images:
            unique_images = []
            for img in images:
                if img not in seen_images:
                    seen_images.add(img)
                    unique_images.append(img)
            if unique_images:
                cell_images[col_idx] = unique_images

    return unique_cells, cell_images


def format_value_with_images(text: str, images: List[str]) -> str:
    """
    将文本和图片信息格式化为value字符串

    Args:
        text: 文本内容
        images: 图片文件名列表

    Returns:
        str: 格式化后的value
    """
    if not images:
        return text

    # 如果有图片，在文本后附加图片标记
    image_markers = [f"[图片: {img}]" for img in images]

    if text:
        return f"{text} {' '.join(image_markers)}"
    else:
        return ' '.join(image_markers)


def extract_table_as_key_value(table, doc_part, output_dir: str, table_idx: int) -> Dict[str, Any]:
    """
    将表格提取为key-value格式
    规则：
    1. 一行多个数据：第一个是key，后面的是value（合并为字符串）
    2. 一行一个数据：这个是key，value在下一行

    Args:
        table: 表格对象
        doc_part: 文档part对象
        output_dir: 图片保存目录
        table_idx: 表格索引

    Returns:
        Dict[str, Any]: key-value格式的表格数据
    """
    result = {}
    rows = list(table.rows)
    i = 0

    while i < len(rows):
        cells, row_images = extract_unique_from_row_with_images(rows[i], doc_part, output_dir, table_idx, i)

        if len(cells) == 0:
            i += 1
            continue

        # 情况1：一行多个数据 - 根据列数判断key和value
        if len(cells) >= 2:
            if len(cells) == 2:
                # 两列：第一列为key，第二列为value
                key = cells[0]
                value_text = cells[1]
                # 获取第二列的图片（列索引1）
                value_images = row_images.get(1, [])
                value = format_value_with_images(value_text, value_images)
            elif len(cells) == 4:
                # 四列：第三列为key，第四列为value
                key = cells[2]
                value_text = cells[3]
                # 获取第四列的图片（列索引3）
                value_images = row_images.get(3, [])
                value = format_value_with_images(value_text, value_images)
            else:
                # 其他情况：第一列为key，后面所有列合并为value
                key = cells[0]
                value_text = ' '.join(cells[1:])
                # 合并所有value列的图片
                value_images = []
                for col_idx in range(1, len(cells)):
                    value_images.extend(row_images.get(col_idx, []))
                value = format_value_with_images(value_text, value_images)
            result[key] = value

        # 情况2：一行一个数据 - 这个是key，value在下一行
        elif len(cells) == 1:
            key = cells[0]
            i += 1
            if i < len(rows):
                next_cells, next_row_images = extract_unique_from_row_with_images(rows[i], doc_part, output_dir, table_idx, i)
                if next_cells:
                    # 如果下一行也是单个数据，作为value
                    if len(next_cells) == 1:
                        value_text = next_cells[0]
                        # 获取下一行第一列的图片
                        value_images = next_row_images.get(0, [])
                        result[key] = format_value_with_images(value_text, value_images)
                    else:
                        # 如果下一行是多个数据，合并作为value
                        value_text = ' '.join(next_cells)
                        # 合并所有列的图片
                        value_images = []
                        for col_idx in next_row_images:
                            value_images.extend(next_row_images[col_idx])
                        result[key] = format_value_with_images(value_text, value_images)
                else:
                    result[key] = ""
                # 处理完value行后，需要再i += 1跳过value行
                i += 1
            else:
                result[key] = ""
            continue

        i += 1

    return result


def read_docx_raw(file_path: str) -> Dict[str, Any]:
    """
    读取Word文档内容，返回原始表格数据（key-value格式）

    Args:
        file_path: Word文档路径

    Returns:
        dict: 包含段落和原始表格数据
    """
    global _image_rId_map

    # 重置全局图片映射，避免多次调用时累积
    _image_rId_map = {}

    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': f'文件不存在: {file_path}',
                'file_path': file_path,
                'file_name': '',
                'paragraphs': [],
                'images': [],
                'tables': [],
                'images_dir': ''
            }

        doc = Document(file_path)

        # 创建图片保存目录（与原文档同路径）
        doc_dir = os.path.dirname(os.path.abspath(file_path))
        doc_name = os.path.splitext(os.path.basename(file_path))[0]
        images_dir = os.path.join(doc_dir, f"{doc_name}_images")

        # 如果目录不存在则创建
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)

        result = {
            'success': True,
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'images_dir': images_dir,
            'paragraphs': [],
            'images': [],  # 新增：存储所有图片信息
            'tables': []
        }

        # 提取段落（包括文本和图片）
        for para_idx, para in enumerate(doc.paragraphs):
            text = clean_text(para.text)

            # 提取段落中的图片
            para_images = []
            if has_image_in_paragraph(para):
                para_images = extract_images_from_paragraph(para, doc.part, images_dir, para_idx + 1)
                if para_images:
                    result['images'].extend(para_images)

            # 如果有文本或图片，添加到段落列表
            if text or para_images:
                para_info = {
                    'index': para_idx + 1,
                    'text': text,
                    'images': para_images
                }
                result['paragraphs'].append(para_info)

        # 提取表格（原始key-value格式，支持图片）
        for table_idx, table in enumerate(doc.tables):
            table_data = extract_table_as_key_value(table, doc.part, images_dir, table_idx + 1)
            if table_data:  # 只添加非空表格
                result['tables'].append({
                    'index': table_idx + 1,
                    'data': table_data
                })

        return result

    except Exception as e:
        import traceback
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc(),
            'file_path': file_path,
            'file_name': os.path.basename(file_path) if file_path else '',
            'images_dir': '',
            'paragraphs': [],
            'images': [],
            'tables': []
        }


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python read_docx_improved.py <docx文件路径> [输出JSON路径]", file=sys.stderr)
        print("\n功能：", file=sys.stderr)
        print("  - 提取Word文档中的原始表格数据", file=sys.stderr)
        print("  - 返回key-value格式的表格内容", file=sys.stderr)
        print("  - 支持提取表格中的图片", file=sys.stderr)
        print("  - 图片保存在与文档同目录的 {文档名}_images 文件夹中", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    # 执行提取
    result = read_docx_raw(file_path)

    # 输出JSON
    json_output = json.dumps(result, ensure_ascii=False, indent=2)

    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(json_output)
        print(f"结果已保存到: {output_path}", file=sys.stderr)

    # 输出到stdout
    print(json_output)


if __name__ == '__main__':
    main()
